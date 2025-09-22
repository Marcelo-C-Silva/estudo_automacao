from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from urllib.parse import urljoin, urlparse
from PIL import Image
import time
import requests

sites = [
    "https://exemplo-site1.com/",
    "https://exemplo-site2.com/",
    "https://exemplo-site3.com/"
]

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5",
}

chrome_options = Options()
chrome_options.add_argument("--headless") 
chrome_options.add_argument("--no-sandbox")
chrome_options.add_argument("--disable-dev-shm-usage")

driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)

document = Document()

for site in sites:
    document.add_paragraph(f"🔗 Site: {site}")
    try:
        driver.get(site)
        time.sleep(3) 

        soup = BeautifulSoup(driver.page_source, "html.parser")

        imagens_encontradas = []

        for img in soup.find_all("img"):
            src = img.get("src", "")
            alt = img.get("alt", "")
            if any(keyword in src.lower() or keyword in alt.lower() 
                   for keyword in ["logo", "brasao", "marca"]):
                full_url = urljoin(site, src)
                parsed = urlparse(full_url)
                normalized_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
                imagens_encontradas.append((normalized_url, img))

        if imagens_encontradas:
            document.add_paragraph("🟢 Imagens encontradas no site:")
            for img_url, _ in imagens_encontradas:
                document.add_paragraph(f"- {img_url}")
        else:
            document.add_paragraph("⚠ Nenhum logo/brasão/marca encontrado")

        for img_url, img_tag in imagens_encontradas:
            try:
                ext = img_url.split(".")[-1].lower()
                if ext == "svg":
                    document.add_paragraph(f"⚠ Ignorada SVG (Word não suporta): {img_url}")
                    continue

                img_response = requests.get(img_url, headers=HEADERS, timeout=10)
                img_response.raise_for_status()
                img_bytes = BytesIO(img_response.content)

                with Image.open(img_bytes) as pil_img:
                    width, height = pil_img.size
                img_bytes.seek(0)

                try:
                    document.add_picture(img_bytes)
                except Exception:
                    document.add_paragraph("⚠ Não foi possível inserir a imagem no Word (formato incompatível)")

                info = f"Imagem: {img_url}\n➡ Tamanho real: {width}x{height}px"
                if img_tag.get("width") or img_tag.get("height"):
                    info += f" | HTML: {img_tag.get('width')}x{img_tag.get('height')}"
                document.add_paragraph(info)

            except requests.exceptions.HTTPError as e:
                document.add_paragraph(f"❌ Não permitido (HTTP {e.response.status_code}): {img_url}")
            except Exception as e:
                document.add_paragraph(f"❌ Erro ao baixar {img_url}: {e}")

    except Exception as e:
        document.add_paragraph(f"❌ Erro no site {site}: {e}")

driver.quit()
document.save("logos_brasoes_marcas_sites.docx")
print("✅ Arquivo 'logos_brasoes_marcas_sites.docx' criado com sucesso!")
